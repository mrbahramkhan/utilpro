#!/usr/bin/env python3
"""
UtilPro PDF Processor - Server-side PDF operations
Usage: python3 pdf_processor.py <action> <args>
"""
import sys, os, json, base64
import pymupdf as fitz
from pathlib import Path

def pdf_info(path):
    doc = fitz.open(path)
    info = {
        "pages": doc.page_count,
        "title": doc.metadata.get("title",""),
        "author": doc.metadata.get("author",""),
        "size_kb": round(os.path.getsize(path)/1024, 1),
        "encrypted": doc.is_encrypted,
    }
    doc.close()
    return info

def merge_pdfs(paths, output):
    result = fitz.open()
    for p in paths:
        doc = fitz.open(p)
        result.insert_pdf(doc)
        doc.close()
    result.save(output)
    result.close()
    return {"success": True, "pages": result.page_count if not result.is_closed else "saved"}

def split_pdf(path, pages, output_dir):
    doc = fitz.open(path)
    outputs = []
    for i, page_num in enumerate(pages):
        out = fitz.open()
        out.insert_pdf(doc, from_page=page_num-1, to_page=page_num-1)
        out_path = os.path.join(output_dir, f"page_{page_num}.pdf")
        out.save(out_path)
        out.close()
        outputs.append(out_path)
    doc.close()
    return {"files": outputs}

def compress_pdf(path, output, quality=50):
    doc = fitz.open(path)
    for page in doc:
        images = page.get_images()
        for img in images:
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            # Re-compress images
    doc.save(output, garbage=4, deflate=True, clean=True)
    orig = os.path.getsize(path)
    comp = os.path.getsize(output)
    doc.close()
    return {"original_kb": round(orig/1024,1), "compressed_kb": round(comp/1024,1), 
            "saved_pct": round((1 - comp/orig)*100, 1)}

def pdf_to_images(path, output_dir, dpi=150, fmt="png"):
    doc = fitz.open(path)
    files = []
    mat = fitz.Matrix(dpi/72, dpi/72)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat, alpha=False)
        out_path = os.path.join(output_dir, f"page_{i+1}.{fmt}")
        if fmt == "jpg":
            pix.save(out_path, output="jpeg", jpg_quality=90)
        else:
            pix.save(out_path)
        files.append(out_path)
    doc.close()
    return {"files": files, "count": len(files)}

def pdf_to_text(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text() + "\n\n"
    doc.close()
    return {"text": text, "chars": len(text)}

def pdf_to_excel(path, output):
    import openpyxl
    doc = fitz.open(path)
    wb = openpyxl.Workbook()
    for i, page in enumerate(doc):
        ws = wb.create_sheet(f"Page {i+1}") if i > 0 else wb.active
        ws.title = f"Page {i+1}"
        tables = page.find_tables()
        if tables and tables.tables:
            for table in tables.tables:
                df = table.to_pandas()
                for r_idx, row in enumerate(df.itertuples(index=False), 1):
                    for c_idx, val in enumerate(row, 1):
                        ws.cell(row=r_idx, column=c_idx, value=str(val) if val else "")
        else:
            # Fallback: extract text line by line
            text = page.get_text()
            for r_idx, line in enumerate(text.split('\n'), 1):
                ws.cell(row=r_idx, column=1, value=line)
    wb.save(output)
    doc.close()
    return {"success": True, "sheets": len(wb.worksheets)}

def pdf_to_word(path, output):
    from docx import Document
    doc_pdf = fitz.open(path)
    doc_word = Document()
    doc_word.add_heading('Converted from PDF', 0)
    for i, page in enumerate(doc_pdf):
        doc_word.add_heading(f'Page {i+1}', level=2)
        blocks = page.get_text("blocks")
        for block in blocks:
            text = block[4].strip()
            if text:
                doc_word.add_paragraph(text)
        if i < doc_pdf.page_count - 1:
            doc_word.add_page_break()
    doc_pdf.close()
    doc_word.save(output)
    return {"success": True, "pages": doc_pdf.page_count}

def images_to_pdf(image_paths, output):
    doc = fitz.open()
    for img_path in image_paths:
        img_doc = fitz.open(img_path)
        rect = img_doc[0].rect
        pdf_bytes = img_doc.convert_to_pdf()
        img_doc.close()
        img_pdf = fitz.open("pdf", pdf_bytes)
        doc.insert_pdf(img_pdf)
    doc.save(output)
    doc.close()
    return {"success": True, "pages": len(image_paths)}

def rotate_pdf(path, output, degrees=90, page_nums=None):
    doc = fitz.open(path)
    pages = page_nums if page_nums else list(range(doc.page_count))
    for i in pages:
        if i < doc.page_count:
            doc[i].set_rotation(degrees)
    doc.save(output)
    doc.close()
    return {"success": True}

def add_watermark(path, output, text="CONFIDENTIAL"):
    doc = fitz.open(path)
    for page in doc:
        rect = page.rect
        # Insert diagonal watermark text
        page.insert_text(
            (rect.width/4, rect.height/2),
            text, fontsize=60, color=(0.8,0.8,0.8),
            rotate=45, overlay=True
        )
    doc.save(output)
    doc.close()
    return {"success": True}

def add_page_numbers(path, output):
    doc = fitz.open(path)
    for i, page in enumerate(doc):
        rect = page.rect
        page.insert_text(
            (rect.width/2 - 20, rect.height - 20),
            f"- {i+1} -", fontsize=10, color=(0,0,0)
        )
    doc.save(output)
    doc.close()
    return {"success": True}

def protect_pdf(path, output, password):
    doc = fitz.open(path)
    perm = fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY
    doc.save(output, encryption=fitz.PDF_ENCRYPT_AES_256,
             user_pw=password, owner_pw=password+"owner", permissions=perm)
    doc.close()
    return {"success": True}

def unlock_pdf(path, output, password):
    doc = fitz.open(path)
    if doc.is_encrypted:
        if not doc.authenticate(password):
            return {"error": "Wrong password"}
    doc.save(output)
    doc.close()
    return {"success": True}

if __name__ == "__main__":
    action = sys.argv[1] if len(sys.argv) > 1 else "info"
    print(json.dumps({"status": "PDF Processor ready", "action": action}))
